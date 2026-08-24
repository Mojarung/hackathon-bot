"""Pulling readable text out of attachments.

The vision model only takes images, so a PDF of the rules would otherwise be
stored and ignored. Text is extracted here and fed to the model alongside the
message, which is both cheaper and more accurate than rendering pages as images.

Scanned PDFs carry no text layer; those are reported as such so the caller can
tell the user to send a screenshot instead of silently extracting nothing.
"""

from __future__ import annotations

import asyncio
import csv
import io
import json
import logging
from dataclasses import dataclass

log = logging.getLogger(__name__)

MAX_CHARS = 20_000
_MIN_USEFUL_CHARS = 80

PLAIN_SUFFIXES = (".txt", ".md", ".markdown", ".csv", ".json", ".log", ".yml", ".yaml", ".rst")
PLAIN_MIMES = ("text/", "application/json", "application/x-yaml")


@dataclass(frozen=True, slots=True)
class Extracted:
    text: str
    kind: str                 # plain | pdf | docx
    truncated: bool = False
    needs_ocr: bool = False   # a PDF with no text layer


def _suffix(file_name: str) -> str:
    _, _, ext = file_name.rpartition(".")
    return f".{ext.casefold()}" if ext else ""


def is_readable(file_name: str, mime: str | None) -> bool:
    """Whether `extract` has any chance with this attachment."""
    suffix = _suffix(file_name)
    mime = (mime or "").casefold()
    if suffix in PLAIN_SUFFIXES or any(mime.startswith(m) for m in PLAIN_MIMES):
        return True
    if suffix == ".pdf" or mime == "application/pdf":
        return True
    return suffix == ".docx" or "wordprocessingml" in mime


def _decode(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _plain(payload: bytes, suffix: str) -> str:
    text = _decode(payload)
    if suffix == ".json":
        try:
            return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return text
    if suffix == ".csv":
        try:
            rows = list(csv.reader(io.StringIO(text)))
            return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows[:200])
        except csv.Error:
            return text
    return text


def _pdf(payload: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(payload))
    chunks: list[str] = []
    for page in reader.pages[:40]:  # a rulebook longer than this is not worth the tokens
        try:
            chunks.append(page.extract_text() or "")
        except Exception as exc:
            log.debug("pdf page failed: %s", exc)
    return "\n\n".join(c.strip() for c in chunks if c.strip())


def _docx(payload: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(payload))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_sync(file_name: str, mime: str | None, payload: bytes) -> Extracted | None:
    suffix = _suffix(file_name)
    mime = (mime or "").casefold()

    if suffix == ".pdf" or mime == "application/pdf":
        text, kind = _pdf(payload), "pdf"
        if len(text.strip()) < _MIN_USEFUL_CHARS:
            return Extracted(text="", kind=kind, needs_ocr=True)
    elif suffix == ".docx" or "wordprocessingml" in mime:
        text, kind = _docx(payload), "docx"
    elif suffix in PLAIN_SUFFIXES or any(mime.startswith(m) for m in PLAIN_MIMES):
        text, kind = _plain(payload, suffix), "plain"
    else:
        return None

    text = text.strip()
    if not text:
        return None
    truncated = len(text) > MAX_CHARS
    return Extracted(text=text[:MAX_CHARS], kind=kind, truncated=truncated)


async def extract(file_name: str, mime: str | None, payload: bytes) -> Extracted | None:
    """Extract text off the event loop; parsing a big PDF is genuinely slow."""
    try:
        return await asyncio.to_thread(_extract_sync, file_name, mime, payload)
    except Exception as exc:
        log.warning("could not read %s: %s", file_name, exc)
        return None


def as_prompt_block(file_name: str, extracted: Extracted) -> str:
    """Wrap extracted text so the model knows where it came from."""
    header = f"Содержимое файла {file_name}"
    if extracted.truncated:
        header += " (начало, файл длиннее)"
    return f"{header}:\n{extracted.text}"
